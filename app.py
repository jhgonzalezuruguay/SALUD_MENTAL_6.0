import streamlit as st
import pandas as pd
from datetime import datetime
from docx import Document
from io import BytesIO
import matplotlib.pyplot as plt
import hashlib

st.set_page_config(page_title="Salud Mental 6.0 - Sueños", layout="centered")

# --------- INICIALIZACIÓN DE VARIABLES DE SESIÓN ---------
if "usuarios" not in st.session_state:
    st.session_state["usuarios"] = [
        {"username": "admin", "password": hashlib.sha256("admin123".encode()).hexdigest(), "rol": "admin"}
    ]
if "user" not in st.session_state:
    st.session_state["user"] = None
if "dreams" not in st.session_state:
    st.session_state["dreams"] = []
if "reset_form" not in st.session_state:
    st.session_state["reset_form"] = False

# --------- CONTROL DE RERUN (Siempre después de inicialización) ---------
if st.session_state.get("do_rerun", False):
    st.session_state["do_rerun"] = False
    st.rerun()

# --------- CONFIGURACIÓN DE CATEGORÍAS ---------
CATEGORIAS_SUEÑOS = [
    "😌 Bueno", "😱 Pesadilla", "😐 Neutral", "🌀 Confuso",
    "👨‍👩‍👧‍👦 Familiares", "👫 Amigos", "💑 Novia/Novio",
    "🦅 Volar", "🌊 Agua", "🏃‍♂️ Persecución"
]
COLORES = [
    "#A5D6A7", "#EF9A9A", "#FFF59D", "#B3E5FC", "#FFDAB9",
    "#E6E6FA", "#FFB6C1", "#87CEEB", "#00BFFF", "#B0C4DE"
]

# --------- FUNCIONES AUXILIARES DE USUARIO ---------
def hash_password(password: str) -> str:
    return hashlib.sha256(password.encode()).hexdigest()

def get_user(username: str):
    return next((u for u in st.session_state["usuarios"] if u["username"] == username), None)

# --------- INTERFAZ DE AUTENTICACIÓN ---------
def mostrar_login():
    st.title("🔐 Salud Mental 6.0 - Iniciar Sesión")
    tabs = st.tabs(["Iniciar sesión", "Registrarse"])
    with tabs[0]:
        username = st.text_input("Usuario", key="login_user")
        password = st.text_input("Contraseña", type="password", key="login_pass")
        if st.button("Ingresar"):
            user = get_user(username)
            if user and user["password"] == hash_password(password):
                st.session_state["user"] = user
                st.success("¡Bienvenido/a!")
                st.session_state["do_rerun"] = True
            else:
                st.error("Usuario o contraseña incorrectos.")
    with tabs[1]:
        username = st.text_input("Nuevo usuario", key="reg_user")
        password = st.text_input("Nueva contraseña", type="password", key="reg_pass")
        if st.button("Registrarse"):
            if not username or not password:
                st.warning("Completa todos los campos.")
            elif get_user(username):
                st.warning("El nombre de usuario ya existe.")
            else:
                st.session_state["usuarios"].append({
                    "username": username,
                    "password": hash_password(password),
                    "rol": "usuario"
                })
                st.success("Usuario registrado. Ahora puedes iniciar sesión.")
                st.session_state["do_rerun"] = True

# --------- BLOQUE DE AUTENTICACIÓN ---------
if not st.session_state["user"]:
    mostrar_login()
    st.stop()

# --------- VARIABLES DE USUARIO Y ROL ---------
user = st.session_state["user"]
es_admin = user["rol"] == "admin"
st.title(f"Salud Mental 6.0 {'(Admin)' if es_admin else ''} - Sueños")

# --------- FORMULARIO PARA NUEVO SUEÑO ---------
with st.form("dream_form"):
    st.subheader("✍️ Escribe tu sueño")
    # --- Manejo del reseteo del formulario ---
    if st.session_state.reset_form:
        titulo_default = ""
        contenido_default = ""
        categoria_default = CATEGORIAS_SUEÑOS[0]
        st.session_state.reset_form = False
    else:
        titulo_default = st.session_state.get("titulo", "")
        contenido_default = st.session_state.get("contenido", "")
        categoria_default = st.session_state.get("categoria", CATEGORIAS_SUEÑOS[0])

    titulo = st.text_input("Título del sueño", value=titulo_default, key="titulo")
    contenido = st.text_area("Describe tu sueño...", value=contenido_default, key="contenido")
    categoria = st.selectbox(
        "Categoría del sueño",
        CATEGORIAS_SUEÑOS,
        index=CATEGORIAS_SUEÑOS.index(categoria_default),
        key="categoria"
    )
    submitted = st.form_submit_button("Guardar sueño")

    if submitted:
        if not contenido.strip():
            st.error("Por favor, escribe el contenido del sueño.")
        elif not titulo.strip():
            st.error("Por favor, asigna un título a tu sueño.")
        else:
            st.session_state["dreams"].append({
                "fecha": datetime.now().strftime("%Y-%m-%d %H:%M"),
                "usuario": user["username"],
                "titulo": titulo.strip()[:100],
                "contenido": contenido.strip()[:1000],
                "categoria": categoria,
                "eliminado": False
            })
            st.success("🌙 Sueño guardado correctamente.")
            st.session_state.reset_form = True
            st.rerun()

# --------- VISTA DE SUEÑOS Y ACCIONES ---------
df = pd.DataFrame(st.session_state["dreams"])

# Si no hay sueños o columnas, muestra mensaje y evita errores
if df.empty or not all(col in df.columns for col in ["eliminado", "usuario", "fecha"]):
    st.subheader("🗂 Tus sueños" if not es_admin else "🗂 Sueños de todos los usuarios")
    st.info("Aún no hay sueños registrados.")
    df_vista = pd.DataFrame()  # Para evitar errores en estadísticas
else:
    if es_admin:
        df_vista = df[df["eliminado"] == False].sort_values(by="fecha", ascending=False)
    else:
        df_vista = df[(df["eliminado"] == False) & (df["usuario"] == user["username"])].sort_values(by="fecha", ascending=False)
    st.subheader("🗂 Tus sueños" if not es_admin else "🗂 Sueños de todos los usuarios")

    if not df_vista.empty:
        st.dataframe(df_vista[["fecha", "titulo", "categoria"]], use_container_width=True)
        for idx, row in df_vista.iterrows():
            with st.expander(f"{row['fecha']} - {row['titulo']} ({row['categoria']})"):
                st.write(row['contenido'])
                if not es_admin and st.button(f"🗑️ Ocultar este sueño", key=f"del_{idx}"):
                    st.session_state["dreams"][df.index[idx]]["eliminado"] = True
                    st.session_state["do_rerun"] = True

        # Solo el admin puede descargar todos los sueños
        if es_admin:
            csv = df.drop(columns=["eliminado"]).to_csv(index=False).encode('utf-8')
            st.download_button("⬇️ Descargar todos como CSV", csv, "suenos_salud_mental.csv", "text/csv")
            doc = Document()
            doc.add_heading("Sueños - Salud Mental 6.0", 0)
            for row in df.itertuples():
                doc.add_heading(f"{row.titulo} ({row.categoria})", level=2)
                doc.add_paragraph(f"Fecha: {row.fecha}")
                doc.add_paragraph(f"Usuario: {row.usuario}")
                doc.add_paragraph(row.contenido)
                doc.add_paragraph("")
            doc_io = BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            st.download_button(
                "⬇️ Descargar todos como Word",
                doc_io,
                "suenos_salud_mental.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        st.info("Aún no hay sueños registrados.")

# --------- ESTADÍSTICAS DE SUEÑOS (Gráfico de barras horizontales) ---------
st.subheader("📊 Estadística de sueños" if not es_admin else "📊 Estadística global de sueños")
if not df.empty and all(col in df.columns for col in ["eliminado", "usuario", "fecha"]) and not df_vista.empty:
    conteo = df_vista["categoria"].value_counts().reindex(CATEGORIAS_SUEÑOS, fill_value=0)
    fig, ax = plt.subplots()
    ax.barh(conteo.index, conteo.values, color=COLORES)
    ax.set_xlabel("Cantidad de sueños")
    ax.set_ylabel("Categoría")
    ax.set_title("Cantidad de sueños por categoría")
    plt.tight_layout()
    st.pyplot(fig)
    resumen = "\n".join([f"- {cat}: {conteo[cat]}" for cat in CATEGORIAS_SUEÑOS])
    st.markdown(f"**Resumen:**\n{resumen}")
else:
    st.info("Registra sueños para ver estadísticas.")

# --------- CIERRE DE SESIÓN ---------
if st.button("Cerrar sesión"):
    st.session_state["user"] = None
    st.session_state["do_rerun"] = True




"""import streamlit as st
import pandas as pd
from datetime import datetime
import os
import csv
import base64

# ================== CONFIGURACIÓN DE ADMIN ==================
ADMIN_USER_CODE = "16990037"  # Cambia esto por tu código personal de 7 dígitos de administrador
ADMIN_PASSWORD = "16990037"    # Cambia esto por tu clave secreta de administrador

# ================== FUNCIONES ===============================

DIARIO_CSV = "diario_emocional.csv"

def guardar_diario_csv(entry):
    file_exists = os.path.isfile(DIARIO_CSV)
    entry_to_save = entry.copy()
    entry_to_save["emociones"] = ";".join(entry["emociones"]) if entry["emociones"] else ""
    entry_to_save["intensidades"] = ";".join([f"{k}:{v}" for k, v in entry["intensidades"].items()]) if entry["intensidades"] else ""
    entry_to_save["acciones"] = ";".join(entry["acciones"]) if entry["acciones"] else ""
    with open(DIARIO_CSV, mode='a', newline='', encoding='utf-8') as file:
        writer = csv.DictWriter(file, fieldnames=entry_to_save.keys())
        if not file_exists:
            writer.writeheader()
        writer.writerow(entry_to_save)

def cargar_diario_csv():
    if os.path.isfile(DIARIO_CSV):
        try:
            df = pd.read_csv(DIARIO_CSV, dtype={'codigo_usuario':str})
        except Exception:
            df = pd.read_csv(DIARIO_CSV, dtype=str)
        return df
    else:
        return pd.DataFrame(columns=["codigo_usuario","fecha", "emociones", "intensidades", "contexto", "acciones"])

def get_table_download_link(df, filename="diario_emocional.csv"):
    csv_str = df.to_csv(index=False, encoding='utf-8')
    b64 = base64.b64encode(csv_str.encode()).decode()
    href = f'<a href="data:file/csv;base64,{b64}" download="{filename}">Descargar historial como CSV</a>'
    return href

# ================== SESIONES STREAMLIT ======================
if "diary_data" not in st.session_state:
    st.session_state.diary_data = []
if "codigo_usuario" not in st.session_state:
    st.session_state.codigo_usuario = None
if "admin_ok" not in st.session_state:
    st.session_state.admin_ok = False

# ================== TÍTULOS ================================
st.title("🌈 VITAL")
st.title("Asistente de Salud Mental con I.A.")
st.title("📔 Diario Emocional: Check-in")

# =========== INGRESO DE USUARIO (CÓDIGO IDENTIFICADOR) ===========
if st.session_state.codigo_usuario is None:
    st.subheader("🔒 Ingreso de Usuario")
    codigo_input = st.text_input("Por favor, ingresa tu Documento de Identidad, sin puntos ni guiones:", max_chars=8)
    if st.button("Ingresar"):
        if codigo_input.isdigit() and len(codigo_input) == 8:
            st.session_state.codigo_usuario = codigo_input
            st.success("¡Código aceptado! Ahora puedes completar tu diario emocional, oprime 'Ingresar' nuevamente.")
        else:
            st.error("El código debe ser numérico y tener exactamente 8 dígitos.")
    st.stop()

# =========== ENTRADA DIARIO EMOCIONAL ======================
emotions = [
    {"label": "Alegría", "emoji": "😊"},
    {"label": "Ansiedad", "emoji": "😟"},
    {"label": "Tristeza", "emoji": "😢"},
    {"label": "Ira", "emoji": "😠"},
    {"label": "Confusión", "emoji": "😐"},
    {"label": "Gratitud", "emoji": "🙏"},
    {"label": "Miedo", "emoji": "😨"}
]

st.subheader("¿Cómo te sientes hoy?")

selected_emotions = st.multiselect(
    "Selecciona hasta 3 emociones:",
    options=[f"{e['emoji']} {e['label']}" for e in emotions],
    max_selections=3
)

emotion_intensities = {}
for emotion in selected_emotions:
    level = st.slider(f"Intensidad de {emotion} (1-10):", 1, 10, 5)
    emotion_intensities[emotion] = level

context = st.text_area("¿Qué hizo que te sintieras así?", placeholder="Describe brevemente lo que pasó...")

st.markdown("¿Qué hiciste para cuidarte?")
coping_actions = st.multiselect(
    "Selecciona las acciones que tomaste:",
    ["Hablé con alguien", "Medité", "Salí a caminar", "Escuché música", "Escribí", "Nada en particular"]
)

if st.button("💾 Guardar entrada de hoy"):
    entry = {
        "codigo_usuario": st.session_state.codigo_usuario,
        "fecha": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "emociones": selected_emotions,
        "intensidades": emotion_intensities,
        "contexto": context,
        "acciones": coping_actions
    }
    st.session_state.diary_data.append(entry)
    guardar_diario_csv(entry)
    st.success("✔️ Entrada guardada exitosamente.")

# =========== HISTORIAL INDIVIDUAL (SIN DESCARGA) =================
if st.checkbox("📖 Mostrar historial de entradas"):
    df = cargar_diario_csv()
    if not df.empty:
        df_usuario = df[df["codigo_usuario"] == st.session_state.codigo_usuario]
        if not df_usuario.empty:
            st.dataframe(df_usuario)
            #st.info("Solo el administrador puede descargar el historial en CSV.")
        else:
            st.info("Aún no has registrado entradas.")
    else:
        st.info("Aún no has registrado entradas.")

# =========== ACCESO ADMINISTRATIVO (SOLO VISIBLE A ADMIN) =========
if st.session_state.codigo_usuario == ADMIN_USER_CODE:
    st.markdown("---")
    st.subheader("🔑 Acceso administrativo (descarga de datos)")

    if not st.session_state.admin_ok:
        admin_code = st.text_input("Código de administrador:", type="password")
        if st.button("Ingresar como administrador"):
            if admin_code == ADMIN_PASSWORD:
                st.session_state.admin_ok = True
                st.success("Acceso concedido. Puedes descargar los historiales.")
            else:
                st.error("Código incorrecto.")
    else:
        st.success("🟢 Acceso de administrador activo.")

        # Listado de códigos únicos de usuario
        st.markdown("#### Códigos de usuario registrados")
        df = cargar_diario_csv()
        codigos_unicos = sorted(df["codigo_usuario"].unique())
        st.write("Códigos únicos registrados:")
        st.code('\n'.join(codigos_unicos), language="text")

        # Descarga historial individual de cualquier usuario
        st.markdown("#### Descargar historial individual de usuario")
        buscar_codigo = st.text_input("Código identificador de usuario para descargar historial:", max_chars=8, key="descarga_individual")
        if buscar_codigo:
            df_usuario = df[df["codigo_usuario"] == buscar_codigo]
            if not df_usuario.empty:
                st.dataframe(df_usuario)
                st.markdown(get_table_download_link(df_usuario, filename=f"diario_usuario_{buscar_codigo}.csv"), unsafe_allow_html=True)
            else:
                st.info("No hay datos para ese código de usuario.")

        # Descarga historial grupal de todos los usuarios
        st.markdown("#### Descargar historial grupal/completo")
        if not df.empty:
            st.dataframe(df)
            st.markdown(get_table_download_link(df, filename="diario_emocional_completo.csv"), unsafe_allow_html=True)
        else:
            st.info("No hay ingresos registrados aún.")

# =========== ENLACES Y PIE DE PÁGINA =========================
st.markdown("---")
st.subheader("📅 Agendar una consulta con un profesional")
booking_url = "https://forms.gle/MQwofoD14ELSp4Ye7"
st.markdown(f'<a href="{booking_url}" target="_blank"><button style="background-color: #4CAF50; color: white; padding: 10px 24px; border: none; border-radius: 4px; cursor: pointer;">AGENDAR CITA</button></a>', unsafe_allow_html=True)

st.markdown("---")
st.subheader("📋 Registro de Usuario")
registro_url = "https://forms.gle/ZsM2xrWyUUU9ak6z7"
st.markdown(f'<a href="{registro_url}" target="_blank"><button style="background-color: #4CAF50; color: white; padding: 10px 24px; border: none; border-radius: 4px; cursor: pointer;">REGISTRARSE</button></a>', unsafe_allow_html=True)

st.markdown("---")
st.subheader("💬 Enviar Mensaje por WhatsApp")
st.write("Si deseas enviar un mensaje por WhatsApp, haz clic en el siguiente botón:")
whatsapp_url = "https://wa.me/59897304859?text=Hola,%20necesito%20ayuda%20con%20mi%20salud%20mental."
st.markdown(f'<a href="{whatsapp_url}" target="_blank"><button style="background-color: #25D366; color: white; padding: 10px 24px; border: none; border-radius: 4px; cursor: pointer;">Enviar Mensaje</button></a>', unsafe_allow_html=True)

st.markdown("---")
st.write("VITAL LE AGRADECE POR CONFIAR Y USAR NUESTRO SERVICIO ❤️")
st.subheader("⚠️  Por consultas, y/o para participar y brindar tu servicio como profesional de la salud en nuestra app, comunicarse con:")
st.write("Mag. José González Gómez")
st.write("Correo: josehumbertogonzalezgomez@gmail.com")

st.markdown("---")
st.subheader("PROFESIONALES DE LA SALUD REFERENTES:")
st.write("Psic. Natalia Brandl")
st.write("Correo: brandlnatalia@gmail.com")
st.write("Atención presencial en: Germán Barbato 1358. Apto 501, y virtual")
st.write('Tesis: "Adicción a videojuegos como riesgo invisible de suicidio"')
st.write("Ps. Bryan Mora Durán")
st.write("Correo: bryanmoraduran@gmail.com")

st.markdown("---")
st.subheader("SECCIÓN DE TALLERES:")
st.write("Analista en Ciberseguridad Matías Alves")
st.write("Taller de concientización del uso de Redes Sociales (virtual)")
st.write("Correo: matiasalvessarmiento@gmail.com")
st.write("Lic. Guillermo Rodríguez")
st.write("Taller de Informática (virtual)")
st.write("Taller de inglés (virtual)")
st.write("Correo: williamforever2014@gmail.com")

st.markdown("---")
st.write("**Nota:** Esta herramienta proporciona diagnósticos preliminares basados en los síntomas ingresados. No reemplaza una consulta profesional.")"""
